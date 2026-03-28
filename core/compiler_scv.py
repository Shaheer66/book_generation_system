from docx import Document
from docx.shared import Pt
import os

def compile_book_to_docx(book_title, chapters_list):
    """
    Takes a list of chapter dictionaries (sorted) and creates a .docx file.
    chapters_list: [{'title': '...', 'content': '...'}, ...]
    """
    doc = Document()
    
    # Book Title
    title_obj = doc.add_heading(book_title, 0)
    doc.add_page_break()

    for chapter in chapters_list:
        
        doc.add_heading(f"Chapter {chapter['chapter_number']}: {chapter['title']}", level=1)
        
        # Chapter Content
        para = doc.add_paragraph(chapter['content'])
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        doc.add_page_break()

    filename = f"exports/{book_title.replace(' ', '_')}.docx"
    os.makedirs("exports", exist_ok=True)
    doc.save(filename)
    return filename